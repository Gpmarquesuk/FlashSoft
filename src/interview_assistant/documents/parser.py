from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

import pdfplumber
import docx


@dataclass
class ParsedDocument:
    sections: Dict[str, List[str]]

    def to_dict(self) -> dict:
        return {"sections": self.sections}


class DocumentParser:
    """
    Converte currículo e job description em estruturas normalizadas.
    """

    def __init__(self) -> None:
        pass

    @staticmethod
    def _clean_lines(lines: List[str]) -> List[str]:
        cleaned: List[str] = []
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned.append(stripped)
        return cleaned

    def parse_pdf(self, path: Path) -> ParsedDocument:
        sections: Dict[str, List[str]] = {}
        try:
            with pdfplumber.open(str(path)) as pdf:
                all_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception:
            all_text = path.read_text(encoding="utf-8", errors="ignore")
        sections["content"] = self._clean_lines(all_text.splitlines())
        return ParsedDocument(sections=sections)

    def parse_docx(self, path: Path) -> ParsedDocument:
        document = docx.Document(str(path))
        sections: Dict[str, List[str]] = {}
        current_heading = "content"
        sections[current_heading] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if "Heading" in style and text.lower() not in sections:
                current_heading = text.lower()
                sections.setdefault(current_heading, [])
            sections[current_heading].append(text)
        for key in list(sections.keys()):
            sections[key] = self._clean_lines(sections[key])
        return ParsedDocument(sections=sections)

    def parse_resume(self, path: Path) -> dict:
        parsed = self._parse(path)
        return {"resume": parsed.to_dict()}

    def parse_job_description(self, path: Path) -> dict:
        parsed = self._parse(path)
        return {"job_description": parsed.to_dict()}

    def _parse(self, path: Path) -> ParsedDocument:
        path = Path(path)
        if path.suffix.lower() == ".pdf":
            return self.parse_pdf(path)
        if path.suffix.lower() in {".docx"}:
            return self.parse_docx(path)
        raise ValueError(f"Formato de documento não suportado: {path.suffix}")

    def parse_both(self, resume_path: Path, jd_path: Path) -> dict:
        data = {}
        data.update(self.parse_resume(resume_path))
        data.update(self.parse_job_description(jd_path))
        return data

    @staticmethod
    def dump_json(data: dict, output_path: Path) -> None:
        output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
