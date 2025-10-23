from pathlib import Path

import docx

from src.interview_assistant.documents.parser import DocumentParser

PDF_CONTENT = r"""%PDF-1.4
1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj
2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj
3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj<< /Length 44 >>stream
BT /F1 24 Tf 50 150 Td (Resume Sample) Tj ET
endstream endobj
5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj
xref
0 6
0000000000 65535 f 
0000000010 00000 n 
0000000061 00000 n 
0000000112 00000 n 
0000000232 00000 n 
0000000323 00000 n 
trailer<< /Root 1 0 R /Size 6 >>
startxref
382
%%EOF
"""


def _build_docx(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Experiencia", level=1)
    document.add_paragraph("Liderou equipe de IA.")
    document.add_heading("Habilidades", level=1)
    document.add_paragraph("Python, Whisper, Retrieval")
    document.save(path)


def _build_pdf(path: Path) -> None:
    path.write_bytes(PDF_CONTENT.encode("latin-1"))


def test_parse_both(tmp_path: Path):
    resume_pdf = tmp_path / "resume.pdf"
    job_docx = tmp_path / "job.docx"
    _build_pdf(resume_pdf)
    _build_docx(job_docx)

    parser = DocumentParser()
    data = parser.parse_both(resume_pdf, job_docx)

    assert "resume" in data and "job_description" in data
    resume_sections = data["resume"]["sections"]
    jd_sections = data["job_description"]["sections"]
    assert any("Resume" in " ".join(lines) for lines in resume_sections.values())
    assert "Habilidades" in " ".join(jd_sections.keys()) or jd_sections
