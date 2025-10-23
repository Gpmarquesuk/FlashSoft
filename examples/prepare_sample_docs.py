from pathlib import Path
from docx import Document


BASE_DIR = Path(__file__).resolve().parent / "data"


def create_docs() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)

    resume_doc = BASE_DIR / "resume.docx"
    doc = Document()
    doc.add_heading("Resumo Profissional", level=1)
    doc.add_paragraph(
        "Engenheiro de software com foco em assistentes de entrevista em tempo real, pipelines RAG e overlay stealth."
    )
    doc.add_heading("Habilidades", level=1)
    doc.add_paragraph(
        "WASAPI audio capture; Whisper streaming; Retrieval augmentation; UI overlay furtivo; Python."
    )
    doc.save(resume_doc)

    jd_doc = BASE_DIR / "job_description.docx"
    doc = Document()
    doc.add_heading("Requisitos principais", level=1)
    doc.add_paragraph("Experiencia comprovada com captura de audio WASAPI e transcricao Whisper em tempo real.")
    doc.add_paragraph("Construir pipelines RAG de baixa latencia e observabilidade completa.")
    doc.add_paragraph("Criar overlay topmost com hotkeys customizaveis.")
    doc.save(jd_doc)

    pdf = BASE_DIR / "resume.pdf"
    pdf.write_bytes(
        b"%PDF-1.4
%‚„œ”
1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj
"
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj
"
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
"
        b"4 0 obj<< /Length 92 >>stream
"
        b"BT /F1 18 Tf 20 170 Td (Resumo Profissional) Tj 20 140 Td (Projetos com WASAPI,) Tj 20 120 Td (Whisper e RAG.) Tj ET
"
        b"endstream endobj
"
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj
"
        b"xref
0 6
0000000000 65535 f 
0000000010 00000 n 
0000000061 00000 n 
0000000112 00000 n 
0000000245 00000 n 
0000000356 00000 n 
trailer<< /Root 1 0 R /Size 6 >>
startxref
420
%%EOF"
    )

    print(f"Arquivos gerados em {BASE_DIR}")


if __name__ == "__main__":
    create_docs()
